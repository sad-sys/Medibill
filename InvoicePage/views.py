from django.shortcuts import render,redirect
from django.http import HttpResponse,FileResponse
from .forms import LoginForm, surgeryForm, calendarForm, RegisterForm, DetailsForm
from django.views.decorators.csrf import csrf_protect
from django.contrib.auth.decorators import login_required
from django.conf import settings
from django.contrib.auth import login, authenticate
import os
from django.forms import formset_factory
from docx import Document
from datetime import datetime
from .models import CustomUser
import random
import string

def index(request):
    return render(request, "invoicePage/index.html")

@csrf_protect
def loginView(request):
    if request.method == 'POST':
        loginForm = LoginForm(request.POST)
        if loginForm.is_valid():
            email = loginForm.cleaned_data.get('email')
            password = loginForm.cleaned_data.get('password')
            
            try:
                customUser = CustomUser.objects.get(email=email)
                if customUser.check_password(password):
                    login(request, customUser)  # Log the user in
                    userID = customUser.id
                    return redirect('makeInvoice', userID=userID)
                else:
                    print("Invalid credentials")
                    return redirect("loginView")
            except CustomUser.DoesNotExist:
                print("User does not exist")
                return redirect("loginView")
    else:
        loginForm = LoginForm()
    return render(request, "invoicePage/login.html", {"loginForm": loginForm})

@csrf_protect
def register(request):
    if request.method == "POST":
        registerForm = RegisterForm(request.POST)
        if registerForm.is_valid():
            email = registerForm.cleaned_data.get('email')
            password = registerForm.cleaned_data.get('password')
            print (email,password)
            if CustomUser.objects.filter(email = email).exists():
                print ("Not Permitted")
                return redirect(register)
            else:
                user = authenticate(request, email=email, password=password)
                if user is None:
                    user = CustomUser.objects.newUser\
                    (
                        email=email,
                        password=password,
                    )
                    user.save()
                    login(request, user)
                    return redirect('inputDetails', userID = user.id)

    else:
        registerForm = RegisterForm(request.POST)
    return render(request, "invoicePage/register.html", {"registerform": registerForm})

@csrf_protect
def inputDetails(request, userID):
    user = CustomUser.objects.get(id = userID)
    if request.method == "POST":
        detailsForm = DetailsForm(request.POST)
        if detailsForm.is_valid():
            user.sortCode = detailsForm.cleaned_data.get("sortCode")              
            user.phoneNumber = detailsForm.cleaned_data.get("phoneNumber")
            user.address = detailsForm.cleaned_data.get("address")
            user.company = detailsForm.cleaned_data.get("company")
            user.bankDetail = detailsForm.cleaned_data.get("bankDetail")
            print(user, user.sortCode,user.phoneNumber, user.address, user.company, user.bankDetail)
            user.save()
    else:
        detailsForm = DetailsForm()
    return render(request, "invoicePage/inputDetails.html", {"detailsForm":detailsForm, "userID":user.id, "userEmail":user.email})

from django.forms.models import model_to_dict
from datetime import datetime

def makeInvoice(request, userID):
    userDetails = CustomUser.objects.get(id=userID)
    print("Make Invoice Selected")
    allData = []
    CalendarFormSet = formset_factory(calendarForm, extra=1)  # Set extra to 0
    calendar_formset = CalendarFormSet()  # Instantiate formset
    surgery_form = surgeryForm()  # Instantiate surgery form

    if request.method == 'POST':
        surgery_form = surgeryForm(request.POST)
        calendar_formset = CalendarFormSet(request.POST, request.FILES)  # Re-instantiate formset with POST data

        if calendar_formset.is_valid() and surgery_form.is_valid():  # Ensure all forms are valid
            for form in calendar_formset:
                if form.cleaned_data:
                    cleaned_data = form.cleaned_data  # Get cleaned data from each form
                    # Convert date object to string if it exists
                    if 'datesChosen' in cleaned_data and cleaned_data['datesChosen']:
                        cleaned_data['datesChosen'] = cleaned_data['datesChosen'].strftime('%Y-%m-%d')
                    if 'timeChosenStart' in cleaned_data and cleaned_data['timeChosenStart']:
                        cleaned_data['timeChosenStart'] = cleaned_data['timeChosenStart'].strftime('%H:%M:%S')
                    allData.append(cleaned_data)
                    print(allData[-1])

            # Get the target month from the first date if available
            if allData and 'datesChosen' in allData[0]:
                first_date = datetime.strptime(allData[0]['datesChosen'], '%Y-%m-%d')
                target_month = first_date.strftime('%B %Y')
            else:
                target_month = "Unknown Month"

            # Debug print to check what is in cleaned_data
            print("Surgery form cleaned data:", surgery_form.cleaned_data)

            # Ensure invoices field is a list
            if userDetails.invoices is None:
                userDetails.invoices = []

            # Append new data to the invoices field
            userDetails.invoices.append(allData)
            userDetails.save()

            # Create invoice using the cleaned data
            filename = createInvoice(allData, surgery_form.cleaned_data['surgeryChosen'], target_month, userDetails)  # Update to match field name
            return render(request, "invoicePage/makeInvoice.html", {"calendarFormSet": calendar_formset, "surgeryForm": surgery_form, "filename": filename, "userID": userID})
        else:
            print("Formset or surgery form is not valid")

    return render(request, "invoicePage/makeInvoice.html", {"calendarFormSet": calendar_formset, "surgeryForm": surgery_form, "userID": userID})




def download_invoice(request, filename,):
    file_path = os.path.join(settings.BASE_DIR, filename)
    response = FileResponse(open(file_path, 'rb'), as_attachment=True, filename=filename)
    return response


def createInvoice(entries, gp_choice, target_month, userDetails):
    gpName = ["Pondtail Surgery", "Shirley Medical Centre", "Bishopford Road Surgery", "Thornton Road Medical Centre"]
    gpPay = [72.5, 85, 95, 110]

    gpIndex = -1

    for i in range(len(gpName)):
        if gpName[i] == gp_choice:
            gpIndex = i
            gpPay1 = gpPay[i]
            break

    if gpIndex == -1:
        raise ValueError("Invalid surgery choice. Please provide a valid GP pay.")

    document = Document()
    document.add_heading('Invoice', 0)

    today = datetime.today()
    formatted_date = today.strftime("%d/%m/%Y")
    
    topLeft = document.add_paragraph("DATE\n")
    topLeft.add_run(formatted_date)
    topLeft.add_run("\nInvoice for ")
    topLeft.add_run(target_month)

    companyName = str(userDetails.company)
    firstLine = str(userDetails.address)
    postCode = 'CR8 3QP'
    phoneNumber = str(userDetails.phoneNumber)
    email = str(userDetails.email)
    accountNo = str(userDetails.bankDetail)
    sortCode = str(userDetails.sortCode)

    topRight = document.add_paragraph(companyName)
    topRight.add_run("\n")
    topRight.add_run(firstLine)
    topRight.add_run("\n")
    topRight.add_run(postCode)
    topRight.add_run("\n")
    topRight.add_run(phoneNumber)
    topRight.add_run("\n")
    topRight.add_run(email)
    topRight.add_run("\n")
    topRight.add_run(accountNo)
    topRight.add_run("\n")
    topRight.add_run(sortCode)
    topRight.alignment = 2

    bottomLeft = document.add_paragraph("INVOICE TO")
    bottomLeft.add_run("\n")
    bottomLeft.add_run(gp_choice)
    bottomLeft.add_run("\n")

    table1 = document.add_table(rows=1, cols=3, style='Medium Grid 1 Accent 1')

    column = table1.rows[0].cells
    column[0].text = 'Date'
    column[1].text = 'No. of Hours'
    column[2].text = 'Amount'

    total = 0
    for entry in entries:
        row = table1.add_row().cells
        date_str = entry['datesChosen']
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')  # Convert string back to datetime object
        formatted_date = date_obj.strftime('%d/%m/%Y')
        hours = entry['hoursPerDay']
        amount = hours * gpPay1
        row[0].text = formatted_date
        row[1].text = str(hours)
        row[2].text = f"£{amount:.2f}"
        total += amount

    totalRow = table1.add_row().cells
    totalRow[2].text = f"Total: £{total:.2f}"

    # Generate a filename for the invoice with local time
    current_time = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Invoice_{current_time}.docx"
    document.save(filename)
    return filename


