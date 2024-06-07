from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from docx import Document
from datetime import datetime
import sys

class MainWindow(QMainWindow):

    global gpPay
    global gpPay1
    global table1
    global document
    global gpHours

    address = ''
    gpName = ["Pondtail Surgery","Shirley Medical Centre","Bishopford Road Surgery", "Thornton Road Medical Centre"]
    gpPay = [72.5, 85, 95, 110]

    gpChoice = input("Please input which surgery the invoice is for: ")

    gpHours = int(input("Please input the hours worked: "))

    gpIndex = -1

    for i in range(0, len(gpName)):
        if gpName[i] == gpChoice:
            gpIndex = i
            gpPay1 = gpPay[i]
            break

    if gpIndex == -1:
        gpPay1 = int(input("Please input a GP pay: "))

    print("Pay/hour:", gpPay1)



    document = Document()
    document.add_heading('Invoice', 0)

    'Top Left, Invoice; Date, Year/Month,'


    today = datetime.today()
    formatted_date = today.strftime("%d/%m/%Y")

    print("Current date in dd/MM/YYYY format:", formatted_date)

    date = formatted_date
    targetMonth = input("Please enter the month the invoice is for      ")
    topLeft = document.add_paragraph("DATE\n")
    topLeft.add_run(date)
    topLeft.add_run("\nInvoice for ")
    topLeft.add_run(targetMonth)

    'Top Right, Company Name, address, email, account No, sort code'
    'Declare the Variables'
    companyName = 'Sadiber LTD'
    firstLine   = '23 Church Hill'
    postCode    = 'CR8 3QP'
    phoneNumber = '07872614143'
    email       = 'drsophiakhawaja@gmail.com'
    accountNo   =  '14459418'
    sortCode    =  '09-01-29'

    topRight = document.add_paragraph(companyName)
    topRight.add_run("\n")
    topRight.add_run(firstLine)
    topRight.add_run("\n")
    topRight.add_run(postCode)
    topRight.add_run("\n")
    topRight.add_run(phoneNumber)
    topRight.add_run("\n")
    topRight.add_run(email)
    topRight.add_run("\n")
    topRight.add_run(accountNo)
    topRight.add_run("\n")
    topRight.add_run(sortCode)

    topRight.alignment = 2

    'Bottom Left, Who is the Invoice too?'
    addressLine1 = gpChoice

    bottomLeft = document.add_paragraph("INVOICE TO")
    bottomLeft.add_run("\n")
    bottomLeft.add_run(addressLine1)
    bottomLeft.add_run("\n")
    #bottomLeft.add_run(addressLine2)


    'Table for the income'
    table1 = document.add_table(rows=1, cols=3, style = 'Medium Grid 1 Accent 1')

    column = table1.rows[0].cells
    column[0].text = 'Date'
    column[1].text = 'No. of Hours'
    column[2].text = 'Amount'

    global data

    data = \
    [
    ]

    def __init__ (self, *args, **kwargs):
        super(MainWindow, self).__init__(*args, **kwargs)

        self.setWindowTitle("Input for invoices")
        self.setGeometry(100,100,600,400)
        self.UiComponents()

        self.show()


    def UiComponents(self):

        def get_date():

            selected_date = calendar.selectedDate().toString('dd/MM/yyyy')

            gpPay = gpHours * gpPay1

            data.append((selected_date, gpHours, gpPay))
            total = 0
            for i in range(len(table1.rows)-1, 0, -1):
                row = table1.rows[i]
                row._element.getparent().remove(row._element)

            for i, x, y in data:
                # Adding a row and then adding data in it.
                row = table1.add_row().cells
                # Converting id to string as table can only take string input
                row[0].text = str(i)
                row[1].text = str(x)
                row[2].text = "£" + str(y)

                total = total + int(y)

                print(i, x, y)


            totalRow = table1.add_row().cells
            totalRow[2].text = "Total: £" + str(total)

            lenData = len(data)

            print("len data is", lenData)
            print(selected_date)
            print(data)
            document.save("GPDATABASET.docx")

        calendar = QCalendarWidget(self)
        calendar.setGeometry(50, 50, 400, 250)
        calendar.selectionChanged.connect(get_date)

App = QApplication(sys.argv)

window = MainWindow()

sys.exit(App.exec())

